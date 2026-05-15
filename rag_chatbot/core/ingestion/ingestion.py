import re
import fitz
import os
import pickle
import hashlib
from pathlib import Path
from llama_index.core import Document, Settings
from llama_index.core.schema import BaseNode
from llama_index.core.node_parser import SentenceSplitter
from dotenv import load_dotenv
from typing import Any, List
from tqdm import tqdm
from ...setting import RAGSettings

load_dotenv()


def _clean_uuid_from_filename(filename: str) -> str:
    """Remove UUID prefix from filename for display in LLM context.
    
    Example: 'fec52f6e9bc7403497b83f743dae7550_Chinh-sach-nghi-phep.docx' 
             -> 'Chinh-sach-nghi-phep.docx'
    """
    if not filename:
        return filename
    # Match UUID patterns: xxxxxxxx-xxxx-xxxx-xxxx-xxxxxxxxxxxx_ or 32 hex chars without dashes followed by _
    uuid_pattern = r'^[a-f0-9]{8}-?[a-f0-9]{4}-?[a-f0-9]{4}-?[a-f0-9]{4}-?[a-f0-9]{12}_'
    cleaned = re.sub(uuid_pattern, '', filename, flags=re.IGNORECASE)
    return cleaned if cleaned else filename


class LocalDataIngestion:
    def __init__(self, setting: RAGSettings | None = None) -> None:
        self._setting = setting or RAGSettings()
        self._node_store = {}
        self._ingested_file = []
        # Set up cache directory for persistent storage
        self._cache_dir = os.path.abspath(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '..', '..', 'data', 'cache'))
        os.makedirs(self._cache_dir, exist_ok=True)
        self._cache_index_file = os.path.join(self._cache_dir, 'cache_index.pkl')
        self._cache_index = self._load_cache_index()

    def _load_cache_index(self) -> dict:
        """Load the cache index that tracks file hashes"""
        if os.path.exists(self._cache_index_file):
            try:
                with open(self._cache_index_file, 'rb') as f:
                    return pickle.load(f)
            except Exception as e:
                print(f"[CACHE] Warning: Could not load cache index: {e}")
        return {}

    def _save_cache_index(self):
        """Save the cache index"""
        try:
            with open(self._cache_index_file, 'wb') as f:
                pickle.dump(self._cache_index, f)
        except Exception as e:
            print(f"[CACHE] Warning: Could not save cache index: {e}")

    def _get_file_hash(self, file_path: str) -> str:
        """Get MD5 hash of file contents to detect changes"""
        hash_md5 = hashlib.md5()
        try:
            with open(file_path, "rb") as f:
                for chunk in iter(lambda: f.read(4096), b""):
                    hash_md5.update(chunk)
            return hash_md5.hexdigest()
        except Exception:
            return ""

    def _get_cache_path(self, file_name: str) -> str:
        """Get the cache file path for a document"""
        safe_name = file_name.replace('/', '_').replace('\\', '_')
        return os.path.join(self._cache_dir, f"{safe_name}.pkl")

    def _clean_nodes_metadata(self, nodes: List[BaseNode], file_name: str) -> List[BaseNode]:
        """Clean UUID prefixes from file_name metadata in cached nodes."""
        display_name = _clean_uuid_from_filename(file_name)
        for node in nodes:
            if hasattr(node, 'metadata') and node.metadata:
                if 'file_name' in node.metadata:
                    node.metadata['file_name'] = display_name
        return nodes

    def _load_cached_nodes(self, file_name: str, file_path: str) -> List[BaseNode] | None:
        """Load cached nodes if they exist and file hasn't changed"""
        cache_path = self._get_cache_path(file_name)
        
        if not os.path.exists(cache_path):
            return None
        
        # Check if file has changed since caching
        current_hash = self._get_file_hash(file_path)
        cached_hash = self._cache_index.get(file_name, {}).get('hash')
        
        if current_hash != cached_hash:
            print(f"[CACHE] File changed, reprocessing: {file_name}")
            return None
        
        try:
            with open(cache_path, 'rb') as f:
                nodes = pickle.load(f)
                # Clean UUID prefixes from cached metadata for backward compatibility
                nodes = self._clean_nodes_metadata(nodes, file_name)
                print(f"[CACHE] Loaded {len(nodes)} cached nodes for: {file_name}")
                return nodes
        except Exception as e:
            print(f"[CACHE] Could not load cache for {file_name}: {e}")
            return None

    def _save_cached_nodes(self, file_name: str, file_path: str, nodes: List[BaseNode]):
        """Save nodes to cache"""
        cache_path = self._get_cache_path(file_name)
        
        try:
            with open(cache_path, 'wb') as f:
                pickle.dump(nodes, f)
            
            # Update cache index with file hash
            self._cache_index[file_name] = {
                'hash': self._get_file_hash(file_path),
                'node_count': len(nodes)
            }
            self._save_cache_index()
            print(f"[CACHE] Saved {len(nodes)} nodes to cache for: {file_name}")
        except Exception as e:
            print(f"[CACHE] Could not save cache for {file_name}: {e}")

    def _filter_text(self, text):
        # Define the regex pattern.
        pattern = r'[a-zA-Z0-9 \u00C0-\u01B0\u1EA0-\u1EF9`~!@#$%^&*()_\-+=\[\]{}|\\;:\'",.<>/?]+'
        matches = re.findall(pattern, text)
        # Join all matched substrings into a single string
        filtered_text = " ".join(matches)
        # Normalize the text by removing extra whitespaces
        normalized_text = re.sub(r"\s+", " ", filtered_text.strip())

        return normalized_text
    
    def _read_pdf(self, file_path: str) -> list:
        """Read text from PDF file, returns list of (page_num, text) tuples"""
        document = fitz.open(file_path)
        pages_data = []
        
        for page_num, page in enumerate(document, start=1):
            page_text = page.get_text("text")
            page_text = self._filter_text(page_text)
            pages_data.append((page_num, page_text))
        
        document.close()
        
        # Check if we got any text
        total_text = " ".join([text for _, text in pages_data])
        if len(total_text.strip()) < 100:  # Very little text extracted
            print(f"⚠️  Warning: PDF appears to be image-based. Attempting OCR...")
            try:
                # Try OCR if available
                import pytesseract
                from pdf2image import convert_from_path
                from PIL import Image
                
                # Convert PDF pages to images
                images = convert_from_path(file_path)
                pages_data = []
                
                for i, image in enumerate(images, start=1):
                    print(f"   OCR processing page {i}/{len(images)}...")
                    page_text = pytesseract.image_to_string(image)
                    page_text = self._filter_text(page_text)
                    pages_data.append((i, page_text))
                
                print(f"   ✓ OCR completed: {sum(len(t) for _, t in pages_data)} characters extracted")
                
            except ImportError:
                print(f"   ✗ OCR not available. Please install Tesseract OCR.")
                print(f"   See OCR_SETUP.md for installation instructions.")
                print(f"   Returning empty data for this PDF.")
                return []
            except Exception as e:
                print(f"   ✗ OCR failed: {e}")
                print(f"   Returning empty data for this PDF.")
                return []
        
        return pages_data
    
    def _read_txt(self, file_path: str) -> str:
        """Read text from TXT or Markdown file"""
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return self._filter_text(text)
    
    def _read_docx(self, file_path: str) -> str:
        """Read text from DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            all_text = ""
            for para in doc.paragraphs:
                all_text += " " + para.text
            return self._filter_text(all_text.strip())
        except ImportError:
            raise ImportError("python-docx is required to read DOCX files. Install with: pip install python-docx")
    
    def _read_file(self, file_path: str) -> str:
        """Read text from file based on extension"""
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self._read_pdf(file_path)
        elif ext in ['.txt', '.md', '.markdown']:
            return self._read_txt(file_path)
        elif ext == '.docx':
            return self._read_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def store_nodes(
        self,
        input_files: list[str],
        embed_nodes: bool = True,
        embed_model: Any | None = None,
    ) -> List[BaseNode]:
        return_nodes = []
        self._ingested_file = []
        if len(input_files) == 0:
            return return_nodes
        splitter = SentenceSplitter.from_defaults(
            chunk_size=self._setting.ingestion.chunk_size,
            chunk_overlap=self._setting.ingestion.chunk_overlap,
            paragraph_separator=self._setting.ingestion.paragraph_sep,
            secondary_chunking_regex=self._setting.ingestion.chunking_regex,
        )
        if embed_nodes:
            Settings.embed_model = embed_model or Settings.embed_model
        
        cached_count = 0
        processed_count = 0
        
        for input_file in tqdm(input_files, desc="Loading documents"):
            file_name = input_file.strip().split("/")[-1]
            # Also handle Windows paths
            if "\\" in file_name:
                file_name = file_name.split("\\")[-1]
            
            self._ingested_file.append(file_name)
            
            # First check in-memory store
            if file_name in self._node_store:
                return_nodes.extend(self._node_store[file_name])
                cached_count += 1
                continue
            
            # Then check persistent cache
            cached_nodes = self._load_cached_nodes(file_name, input_file)
            if cached_nodes is not None:
                self._node_store[file_name] = cached_nodes
                return_nodes.extend(cached_nodes)
                cached_count += 1
                continue
            
            # Process the file if not cached
            processed_count += 1
            try:
                # Check file type
                ext = Path(input_file).suffix.lower()
                
                # Clean UUID prefix from filename for display in LLM context
                display_name = _clean_uuid_from_filename(file_name)
                
                if ext == '.pdf':
                    # Read PDF with page information
                    pages_data = self._read_pdf(input_file)
                    
                    if not pages_data:
                        print(f"No content extracted from {file_name}")
                        continue
                    
                    # Process each page separately to maintain page boundaries
                    nodes = []
                    for page_num, page_text in pages_data:
                        if page_text.strip():  # Only add pages with content
                            doc = Document(
                                text=page_text,
                                metadata={
                                    "file_name": display_name,
                                    "page_label": str(page_num),
                                },
                            )
                            # Split this page into chunks
                            page_nodes = splitter([doc], show_progress=False)
                            # Ensure all chunks from this page have the correct page_label
                            for node in page_nodes:
                                node.metadata["page_label"] = str(page_num)
                            nodes.extend(page_nodes)
                else:
                    # Read other file types (returns string)
                    all_text = self._read_file(input_file)
                    
                    document = Document(
                        text=all_text,
                        metadata={
                            "file_name": display_name,
                        },
                    )
                    nodes = splitter([document], show_progress=False)
                
                # Embed nodes if needed
                if embed_nodes and nodes:
                    nodes = Settings.embed_model(nodes, show_progress=False)
                
                # Store in memory and persist to cache
                self._node_store[file_name] = nodes
                self._save_cached_nodes(file_name, input_file, nodes)
                return_nodes.extend(nodes)
                
            except Exception as e:
                print(f"Error processing {file_name}: {e}")
                continue
        
        if cached_count > 0 or processed_count > 0:
            print(f"[CACHE] Summary: {cached_count} documents from cache, {processed_count} newly processed")
        
        return return_nodes

    def reset(self):
        self._node_store = {}
        self._ingested_file = []

    def check_nodes_exist(self):
        return len(self._node_store.values()) > 0

    def get_all_nodes(self):
        return_nodes = []
        for nodes in self._node_store.values():
            return_nodes.extend(nodes)
        return return_nodes

    def get_ingested_nodes(self):
        return_nodes = []
        for file in self._ingested_file:
            return_nodes.extend(self._node_store[file])
        return return_nodes

    def get_nodes_for_files(self, file_names: list[str] | None) -> tuple[list[BaseNode], list[str]]:
        """Return nodes for the provided file names along with any missing files."""
        selected_nodes: list[BaseNode] = []
        missing_files: list[str] = []

        if not file_names:
            return selected_nodes, missing_files

        for file_name in file_names:
            if not file_name:
                continue

            normalized = Path(file_name).name
            nodes = self._node_store.get(normalized)

            if nodes:
                selected_nodes.extend(nodes)
            else:
                missing_files.append(normalized)

        return selected_nodes, missing_files
